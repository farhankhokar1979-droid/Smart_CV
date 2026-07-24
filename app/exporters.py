"""Resume export: PDF (via WeasyPrint, rendering the same HTML used for
preview) and DOCX (built natively from section data via python-docx).
Fully offline — no external services involved.
"""
import io
import os

from flask import request
from weasyprint import HTML
import docx
from docx.shared import Pt, RGBColor

from app.resume_renderer import render_resume, build_render_context


def export_pdf(resume):
    """Render the resume's chosen template to HTML, then rasterize to PDF bytes."""
    body_html = render_resume(resume)

    # Bundled, offline-safe fonts — the browser preview loads Inter / Plus Jakarta
    # Sans from Google Fonts over the internet, but WeasyPrint has no internet
    # access during export, so it silently substitutes a system default instead.
    # We embed real local copies via file:// paths (fast, no server round-trip
    # needed) so the PDF's typography matches the on-screen preview.
    fonts_dir = os.path.join(os.path.dirname(__file__), "static", "fonts")
    inter_path = "file://" + os.path.join(fonts_dir, "Inter-Regular.ttf").replace("\\", "/")
    jakarta_path = "file://" + os.path.join(fonts_dir, "PlusJakartaSans.ttf").replace("\\", "/")

    full_html = f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<style>
  @page {{ size: A4; margin: 0; }}
  body {{ margin: 0; }}
  @font-face {{
    font-family: 'Inter';
    src: url('{inter_path}') format('truetype');
    font-weight: 100 900;
  }}
  @font-face {{
    font-family: 'Plus Jakarta Sans';
    src: url('{jakarta_path}') format('truetype');
    font-weight: 200 800;
  }}
</style>
</head><body>{body_html}</body></html>"""

    # base_url lets WeasyPrint resolve relative paths like /static/uploads/... —
    # without it, uploaded profile pictures silently fail to render in the PDF.
    base_url = request.url_root if request else None
    pdf_bytes = HTML(string=full_html, base_url=base_url).write_pdf()
    return pdf_bytes


def export_docx(resume):
    """Build a clean, ATS-friendly DOCX directly from section data (not the
    HTML template — DOCX export intentionally stays plain/parseable)."""
    context = build_render_context(resume)
    style = context["style"]
    sections = context["sections"]

    accent_hex = style["primary_color"].lstrip("#")
    accent_rgb = RGBColor(*(int(accent_hex[i:i + 2], 16) for i in (0, 2, 4))) if len(accent_hex) == 6 else RGBColor(0x6D, 0x28, 0xD9)

    document = docx.Document()

    name = resume.owner.full_name if resume.owner else "Your Name"
    heading = document.add_heading(name, level=1)
    heading.runs[0].font.color.rgb = accent_rgb

    if resume.job_title:
        sub = document.add_paragraph(resume.job_title)
        sub.runs[0].font.size = Pt(12)
        sub.runs[0].font.italic = True

    for section in sections:
        h = document.add_heading(section["label"], level=2)
        if h.runs:
            h.runs[0].font.color.rgb = accent_rgb
            h.runs[0].font.size = Pt(13)

        for entry in section["entries"]:
            values = [v for v in entry.values() if v]
            if not values:
                continue
            p = document.add_paragraph()
            p.add_run(" — ".join(values))

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()
